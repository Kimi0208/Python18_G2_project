from datetime import datetime
from django.http import JsonResponse
from django.shortcuts import reverse, get_object_or_404
from webapp.forms import TaskForm, FileForm
from webapp.models import Task, Status, Priority, Type, File, Checklist, Comment, FileSignature
from django.views.generic import ListView, CreateView, UpdateView, DeleteView, DetailView
from docxtpl import DocxTemplate
from shutil import copyfile
from webapp.views.mail_send import send_email_notification
from django.db.models import ForeignKey
from accounts.models import DefUser, Department
from docx import Document
from docx.shared import Inches


def get_user_info(user_object):
    if user_object.patronymic:
        user_info = {
            'id': user_object.id,
            'first_name': user_object.first_name,
            'last_name': user_object.last_name,
            'patronymic': user_object.patronymic
        }
    else:
        user_info = {
            'id': user_object.id,
            'first_name': user_object.first_name,
            'last_name': user_object.last_name
        }
    return user_info


def get_user_initials(user_object):
    if user_object.patronymic:
        author = (f'{user_object.last_name.capitalize()} {user_object.first_name[0].capitalize()}. '
                  f'{user_object.patronymic[0].capitalize()}.')
    else:
        author = f'{user_object.last_name.capitalize()} {user_object.first_name[0].capitalize()}.'
    return author


class TaskListView(ListView):
    model = Task
    template_name = 'index.html'
    context_object_name = 'tasks'

    def get_queryset(self):
        return Task.objects.filter(destination_to_user=self.request.user.pk)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        users = DefUser.objects.all().order_by('username')
        departments = Department.objects.all().order_by('name')
        checklists = Checklist.objects.all()
        context['users'] = users
        context['departments'] = departments
        context['checklists'] = checklists
        return context


def get_tasks_from_id(request, pk):
    tasks = ''
    if 'tasks/user' in request.path:
        tasks = Task.objects.filter(destination_to_user=pk)
    elif 'tasks/department' in request.path:
        tasks = Task.objects.filter(destination_to_department=pk)
    tasks_list = []
    for task in tasks:
        task_data = {
            'id': task.pk,
            'title': task.title,
            'created_at': task.created_at,
            'deadline': task.deadline,
            'status': task.status.name,
            'priority': task.priority.name,
            'type': task.type.name,
            'author': task.author.username
        }
        tasks_list.append(task_data)
    return JsonResponse({'tasks': tasks_list})


def get_object_from_model(model, value):
    try:
        return model.objects.get(pk=value)
    except model.DoesNotExist:
        return None


def check_is_foreign_key(field, old_value, new_value):
    if isinstance(field, ForeignKey):
        old_obj_name = get_object_from_model(field.related_model, old_value)
        new_obj_name = get_object_from_model(field.related_model, new_value)

        return old_obj_name, new_obj_name
    else:
        return old_value, new_value


def get_files_history(task_pk):
    files_history_list = []
    files_history = File.history.filter(task_id=task_pk)
    for file_history in files_history:
        action = ""
        if file_history.history_type == "+":
            action = "Добавлен файл:"
        elif file_history.history_type == "-":
            action = "Удален файл:"
        author = get_user_initials(file_history.history_user)
        history_info = [(action, file_history.history_date.strftime("%d-%m-%Y %H:%M"),
                         author, file_history.file)]
        files_history_list.append(history_info)
    return files_history_list


def get_comments_history(task_pk):
    comments_history_list = []
    comments_history = Comment.history.filter(task_id=task_pk)
    action = ''
    for comment_history in comments_history:
        if comment_history.history_type == "~":
            old_description = comment_history.prev_record.description
            new_description = comment_history.description
            action = f"Изменен комментарий с {old_description} на {new_description}"
        elif comment_history.history_type == "+":
            action = f"Добавлен комментарий: {comment_history.description}"
        elif comment_history.history_type == "-":
            action = f"Удален комментарий: {comment_history.description}"
        author = get_user_initials(comment_history.history_user)
        history_info = [
            (action, comment_history.history_date.strftime("%d-%m-%Y %H:%M"), author)]
        comments_history_list.append(history_info)
    return comments_history_list


def get_subtask_history(task_pk):
    subtask_history_list = []
    subtasks = Task.objects.filter(parent_task_id=task_pk)
    for subtask in subtasks:
        author = get_user_initials(subtask.author)
        history_info = [('Создана подзадача', subtask.created_at.strftime("%d-%m-%Y %H:%M"), author,
                         subtask.title)]
        subtask_history_list.append(history_info)
    return subtask_history_list


def get_history_task(request, task_pk):
    history_list = []
    task = Task.objects.get(pk=task_pk)
    task_history = list(task.history.all().order_by('history_date'))
    for i in range(1, len(task_history)):
        current_record = task_history[i]
        previous_record = task_history[i - 1]
        delta = current_record.diff_against(previous_record)
        change_date = current_record.history_date.strftime("%d-%m-%Y %H:%M")
        change_user = current_record.history_user
        changes = []
        for change in delta.changes:
            verbose_name = current_record._meta.get_field(change.field).verbose_name
            field = current_record._meta.get_field(change.field)
            if type(change.old) == datetime:
                change.old = change.old.strftime("%d-%m-%Y %H:%M")
            if type(change.new) == datetime:
                change.new = change.new.strftime("%d-%m-%Y %H:%M")
            old, new = check_is_foreign_key(field, change.old, change.new)
            change_author = get_user_initials(change_user)
            change_info = (verbose_name, change_date, change_author, str(old), str(new))
            changes.append(change_info)
        history_list.append(changes)
    history_list.extend(get_files_history(task_pk))
    history_list.extend(get_comments_history(task_pk))
    history_list.extend(get_subtask_history(task_pk))
    task_author = get_user_initials(task.author)
    create_record = [('Создана задача', task.created_at.strftime("%d-%m-%Y %H:%M"), task_author, task.title)]
    history_list.extend([(create_record)])
    if history_list:
        sorted_history = sorted(history_list, key=lambda x: x[0][1], reverse=True)
        return JsonResponse({'history': sorted_history})
    else:
        return JsonResponse({'history': history_list})


def get_task_files(request, task_pk):
    files = File.objects.filter(task=task_pk)
    file_list = []
    for file in files:
        sign_url = ''
        if file.checklist:
            if file.checklist.users.filter(id=request.user.id).exists():
                sign_url = f'sign_file/{file.pk}/'

        file_data = {
            'id': file.id,
            'name': file.file.name,
            'task_id': file.task.id,
            'url': file.file.url,
            'sign_url': sign_url,
            'current_user': request.user.id
        }
        file_list.append(file_data)
    signed_files = []
    files = FileSignature.objects.filter(task_id=task_pk, user=request.user)
    for file in files:
        signed_files.append(file.file.pk)
    return JsonResponse({'files': file_list, "signed_files": signed_files})


def sign_file(request, file_id):
    doc_to_sign = get_object_or_404(File, pk=file_id)

    doc = Document(doc_to_sign.file)

    current_user = request.user
    current_user_id = str(current_user.id)

    if current_user.signature:
        signature_path = current_user.signature.path

    for table in doc.tables:
        for row in table.rows:
            if current_user_id in row.cells[1].text:
                row.cells[1].text = ''
                paragraph = row.cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(signature_path, width=Inches(2))

    doc.save(doc_to_sign.file.path)
    FileSignature.objects.create(file=doc_to_sign, user=current_user, task=doc_to_sign.task)
    return JsonResponse({"success": True, "user_id": current_user_id, "file_id": file_id})


def get_subtasks(object):
    subtasks = Task.objects.filter(parent_task=object).order_by('-id')
    subtasks_list = []
    destination_to = ''
    for subtask in subtasks:
        if subtask.destination_to_department:
            destination_to = subtask.destination_to_department.name
        elif subtask.destination_to_user:
            destination_to = subtask.destination_to_user.username
        subtask_data = {
            'id': subtask.id,
            'title': subtask.title,
            'destination_to': destination_to,
            'author': subtask.author.username,
            'created_at': subtask.created_at,
            'type': subtask.type.name,
            'updated_at': subtask.updated_at,
            'status': subtask.status.name
        }
        subtasks_list.append(subtask_data)
    return subtasks_list


def get_comments(task_pk, user_id):
    comments = Comment.objects.filter(task=task_pk)
    comments_list = []
    for comment in comments:
        author = get_user_info(comment.author)
        comment_data = {
            'id': comment.id,
            'author': author,
            'task': comment.task.id,
            'created_at': comment.created_at,
            'updated_at': comment.updated_at,
            'description': comment.description,
            'user_id': user_id
        }
        comments_list.append(comment_data)
    return comments_list


class TaskView(DetailView):
    model = Task
    permission_required = 'webapp.view_task'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        checklists = Checklist.objects.all()
        context['checklists'] = checklists
        return context

    def render_to_response(self, context, **response_kwargs):
        subtasks_list = get_subtasks(self.object)
        destination_to = ''
        if self.object.parent_task:
            if self.object.parent_task.destination_to_department:
                destination_to = self.object.parent_task.destination_to_department.name
            elif self.object.parent_task.destination_to_user:
                destination_to = self.object.parent_task.destination_to_user.username
            parent_task = {'id': self.object.parent_task.id, 'title': self.object.parent_task.title,
                           'author': self.object.parent_task.author.username,
                           'created_at': self.object.parent_task.created_at, 'destination_to': destination_to,
                           'type': self.object.parent_task.type.name, 'updated_at': self.object.parent_task.updated_at}
        else:
            parent_task = None
        task_data = {
            'id': self.object.pk,
            'title': self.object.title,
            'description': self.object.description,
            'created_at': self.object.created_at,
            'start_date': self.object.start_date,
            'updated_at': self.object.updated_at,
            'done_at': self.object.done_at,
            'deadline': self.object.deadline,
            'status': self.object.status.name,
            'priority': self.object.priority.name,
            'author': self.object.author.username,
            'type': self.object.type.name,
            'parent_task': parent_task,
            'subtasks': subtasks_list,
            'comments': get_comments(self.object.pk, self.request.user.pk)
        }
        return JsonResponse({'task': task_data})


class TaskCreateView(CreateView):
    model = Task
    form_class = TaskForm
    template_name = 'task_proposal_create.html'
    permission_required = 'webapp.add_task'

    def form_invalid(self, form):
        return JsonResponse({'errors': form.errors})

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.author = self.request.user
        destination_to = ''
        if self.object.destination_to_user:
            destination_to = self.object.destination_to_user.username
        elif self.object.destination_to_department:
            destination_to = self.object.destination_to_department.name
        if 'task_pk' in self.kwargs:
            self.object.parent_task = Task.objects.get(pk=self.kwargs['task_pk'])
        self.object.save()
        if self.object.destination_to_user:
            subject = f'CRM: Новая задача #{self.object.pk}  {self.object.title}'
            message = self.object.description
            try:
                send_email_notification(subject, message, self.object.destination_to_user.email)
            except Exception as e:
                print(f"Ошибка при отправке электронного уведомления: {e}")

        task_data = {
            'id': self.object.pk,
            'title': self.object.title,
            'description': self.object.description,
            'created_at': self.object.created_at,
            'start_date': self.object.start_date,
            'updated_at': self.object.updated_at,
            'done_at': self.object.done_at,
            'deadline': self.object.deadline,
            'status': self.object.status.name,
            'priority': self.object.priority.name,
            'author': self.object.author.username,
            'type': self.object.type.name,
            'destination_to': destination_to,
            'subtasks': get_subtasks(self.object.parent_task)
        }
        return JsonResponse(task_data)


class TaskUpdateView(UpdateView):
    model = Task
    form_class = TaskForm
    template_name = 'task_proposal_edit.html'
    permission_required = 'webapp.change_task'

    def form_invalid(self, form):
        return JsonResponse({'errors': form.errors})

    def form_valid(self, form):
        self.object = form.save()
        destination_to = ''
        if self.object.destination_to_user:
            destination_to = self.object.destination_to_user.username
        elif self.object.destination_to_department:
            destination_to = self.object.destination_to_department.name
        if self.object.status.name == 'Выполнена':
            self.object.done_at = datetime.now()
            if self.object.destination_to_user:
                subject = f'CRM: Задача #{self.object.id} выполнена {self.object.title}'
                message = self.object.description
                try:
                    send_email_notification(subject, message, self.object.author.email)
                except Exception as e:
                    print(f"Ошибка при отправке электронного уведомления: {e}")

        task_data = {
            'id': self.object.pk,
            'title': self.object.title,
            'description': self.object.description,
            'start_date': self.object.start_date,
            'updated_at': self.object.updated_at,
            'done_at': self.object.done_at,
            'deadline': self.object.deadline,
            'status': self.object.status.name,
            'priority': self.object.priority.name,
            'type': self.object.type.name,
            'subtasks': get_subtasks(self.object),
            'destination_to': destination_to,
            'created_at': self.object.created_at,
            'author': self.object.author.username
        }
        return JsonResponse(task_data)


class TaskDeleteView(DeleteView):
    model = Task
    template_name = 'task_proposal_delete.html'

    def get_success_url(self):
        return reverse('webapp:index')


def add_subtasks(request, checklist_pk, task_pk):
    main_task = Task.objects.get(pk=task_pk)
    checklist = Checklist.objects.get(pk=checklist_pk)
    users = checklist.users.all()
    title = 'Подпись'
    description = f'Необходима подпись документа в задаче #{task_pk}'
    status = Status.objects.get(pk=1)
    priority = Priority.objects.get(pk=1)
    type = Type.objects.get(pk=1)
    for user in users:
        task = Task.objects.create(author=main_task.author, title=title, description=description, status=status,
                                   priority=priority,
                                   type=type, destination_to_user=user, parent_task=main_task)
        subject = f'CRM: Новая подзадача #{task.id}  {task.title}'
        message = task.description
        try:
            send_email_notification(subject, message, user.email)
        except Exception as e:
            print(f"Ошибка при отправке электронного уведомления: {e}")

    file_count = File.objects.count()
    doc_name = f'Задача{task_pk}_{file_count}'
    base_file_path = 'uploads/user_docs/Шаблон.docx'
    new_file_path = f'uploads/user_docs/{doc_name}.docx'
    copyfile(base_file_path, new_file_path)
    doc = DocxTemplate(new_file_path)

    doc.save(new_file_path)
    for table in doc.tables:
        for row in table.rows:
            if 'Подпись' in row.cells[0].text:
                row.cells[1].text = ''
                paragraph = row.cells[1].paragraphs[0]
                run = paragraph.add_run()
                paragraph.add_run().add_picture(request.user.signature.path, width=Inches(2))
    doc.save(new_file_path)

    created_at = datetime.now().strftime("%d.%m.%Y")
    user_patronymic = ''
    if request.user.patronymic:
        user_patronymic = request.user.patronymic[0] + "."
    user_signature = ''
    if request.user.signature:
        user_signature = request.user.signature
    user = {
        'last_name': request.user.last_name,
        'first_name': request.user.first_name[0] + ".",
        'patronymic': user_patronymic,
        'signature': user_signature
    }
    context = {'title': main_task.title, 'description': main_task.description, 'users': users, 'author': user,
               'created_at': created_at}
    doc.render(context)
    doc.save(new_file_path)
    File.objects.create(user=request.user, task=main_task, file=new_file_path, checklist_id=checklist_pk)
    subtasks = get_subtasks(main_task)
    return JsonResponse({'subtasks': subtasks})


class FileAddView(CreateView):
    model = File
    form_class = FileForm
    template_name = 'file_add.html'

    def form_valid(self, form):
        self.object = form.save(commit=False)
        if self.object.file:
            self.object.user = self.request.user
            self.object.task = Task.objects.get(pk=self.kwargs['task_pk'])
            self.object.save()
            file = {
                'file': self.object.file.name,
            }
            return JsonResponse({'file': file})
        return JsonResponse({'file': None})


class FileDeleteView(DeleteView):
    model = File
    template_name = 'partial/file_delete.html'

    def form_valid(self, form):
        file_id = self.object.id
        self.object.delete()
        return JsonResponse({'file_id': file_id})


def check_new_task(request):
    tasks = Task.objects.filter(destination_to_user=request.user, status_id=1)
    if len(tasks) == 0:
        return JsonResponse({'task_count': 0})
    else:
        return JsonResponse({'task_count': len(tasks)})
